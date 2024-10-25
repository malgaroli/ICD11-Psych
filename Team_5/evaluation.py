import docx
import query_data_all
from query_data_all import query_rag  
import re
from rapidfuzz import fuzz, process



def load_vignettes(filepath):
    """Load all vignettes from a DOCX file, supporting labels with or without trailing letters."""
    doc = docx.Document(filepath)
    vignettes = {}
    current_label = None
    current_vignette = ""

    #"Vignette 1a" or "Vignette 6"
    label_pattern = re.compile(r"Vignette\s+(\d+[a-zA-Z]?)")

    for para in doc.paragraphs:
        text = para.text.strip()

        # Detect vignette label
        label_match = label_pattern.match(text)
        if label_match:
            # Save the previous vignette if exists
            if current_label and current_vignette.strip():
                vignettes[current_label] = current_vignette.strip()

            # Start a new vignette with the current label
            current_label = label_match.group(1)
            current_vignette = ""  # Reset the vignette content
        else:
            # Accumulate paragraphs for the current vignette
            current_vignette += text + "\n"

    # Save the last vignette if any
    if current_label and current_vignette.strip():
        vignettes[current_label] = current_vignette.strip()

    return vignettes

def load_ground_truth(filepath):
    """Load ground truth ICD-11 diagnoses from a DOCX file."""
    doc = docx.Document(filepath)
    ground_truth = {}

    vignette_pattern = re.compile(r"Vignette\s+(\d+[a-zA-Z]?)")
    diagnosis_pattern = re.compile(r"ICD-11:\s*([a-zA-Z0-9\s\-]+)")

    current_label = None

    for para in doc.paragraphs:
        text = para.text.strip()

        label_match = vignette_pattern.match(text)
        if label_match:
            current_label = label_match.group(1) 
            continue  

        if current_label:
            diagnosis_match = diagnosis_pattern.search(text)
            if diagnosis_match:
                diagnosis = diagnosis_match.group(1).strip()  
                ground_truth[current_label] = diagnosis
                current_label = None  

    return ground_truth



def evaluate_model(vignettes, ground_truth):
    """Evaluate model performance by comparing generated diagnoses with ground truth ICD-11 diagnoses."""
    correct_count = 0
    total_count = len(vignettes)

    for label, vignette in vignettes.items():
        print(f"Processing {label}...")

        # Generate diagnosis using the model
        model_diagnosis = query_rag(vignette)  

        # Extract the final diagnosis from the model's response
        final_diagnosis = extract_final_diagnosis(model_diagnosis)
        print(f"Model Diagnosis: {final_diagnosis}")

        # Get the ground truth diagnosis for the current vignette
        ground_truth_diagnosis = ground_truth.get(label, "")
        print(f"Ground Truth: {ground_truth_diagnosis}")

        # Compare using fuzzy matching for flexible comparison
        similarity = fuzz.partial_ratio(final_diagnosis.lower(), ground_truth_diagnosis.lower())
        print(f"Similarity Score: {similarity}")

        if similarity >= 80:  
            correct_count += 1


    accuracy = correct_count / total_count if total_count > 0 else 0
    print(f"Final Accuracy: {accuracy:.2f}")

def extract_final_diagnosis(response):
    """Extract the final diagnosis from the model's response."""
    patterns = [
        r"Final Diagnosis:\s*(.+?)\s*(?:\(|\n|$)", 
        r"diagnosis is\s*(.+?)(?:\n|$)", 
        r"Final Diagnosis:\s*(.+?)(?:\n|$)",
        r"Final Diagnosis:\s*(.+?)\s*(?:\(|\n|$)",
        r"diagnosis is\s*(.+?)(?:\n|$)", 
        r"(?:Final Diagnosis:\s*)+(.+?)(?:\(|\n|$)",
        r"(?:Final Diagnosis:\s*)(.+?)(?:\(|\n|$)",
    ]

    # Try all patterns to find the diagnosis
    for pattern in patterns:
        match = re.search(pattern, response, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "Diagnosis not found"


if __name__ == "__main__":
    vignettes = load_vignettes("/Users/xzh/Desktop/CDS_Capstone/Vignettes_Mood.docx")
    ground_truth = load_ground_truth("/Users/xzh/Desktop/CDS_Capstone/Mood Disorders Vignette Instructions and Descriptions.docx")

    # Evaluate the model's performance
    evaluate_model(vignettes, ground_truth)

